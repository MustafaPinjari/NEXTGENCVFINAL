# resume/views.py
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from .forms import ResumeForm
from .models import Resume
from django.views.decorators.csrf import csrf_exempt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

@login_required
@csrf_exempt  # Use this only for testing; consider using CSRF tokens in production
def create_resume(request):
    if request.method == "POST":
        form = ResumeForm(request.POST, request.FILES)
        if form.is_valid():
            # Save the resume to the database
            resume = form.save(commit=False)
            resume.user = request.user
            resume.save()
            print(resume.about_you)  # Check if this prints the expected value

            # Return the saved data for live preview
            data = {
                "first_name": resume.First_Name,
                "last_name": resume.Last_Name,
                "email": resume.email,
                "phone": resume.phone,
                "address": resume.address,
                "nationality": resume.nationality,
                "summary": resume.summary,
                "about_you": resume.about_you,
                "company_name": resume.company_name,
                "job_title": resume.job_title,
                "work_from": resume.work_from,
                "work_to": resume.work_to,
                "work_description": resume.work_description,
                "institution_name": resume.institution_name,
                "degree": resume.degree,
                "edu_from": resume.edu_from,
                "edu_to": resume.edu_to,
                "edu_description": resume.edu_description,
                "skill_name": resume.skill_name,
                "skill_level": resume.skill_level,
            }
            return JsonResponse(data)
        else:
            return JsonResponse({"error": form.errors}, status=400)

    form = ResumeForm()
    return render(request, "resume/create_resume.html", {"form": form})

@login_required
@csrf_exempt  # Use this only for testing; consider using CSRF tokens in production
def download_resume(request, format):
    if request.method == 'POST':
        try:
            resume = Resume.objects.get(user=request.user)

            if format == "pdf":
                return generate_pdf(resume)
            elif format == "docx":
                return generate_docx(resume)
            elif format == "jpeg":
                return generate_jpeg(resume)
            else:
                return HttpResponse("Unsupported format", status=400)

        except Resume.DoesNotExist:
            return HttpResponse("Resume not found", status=404)

# Helper function to generate PDF
def generate_pdf(resume):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)

    # Use the correct attribute names
    p.drawString(100, 750, f"Name: {resume.First_Name} {resume.Last_Name}")
    p.drawString(100, 730, f"Email: {resume.email}")
    p.drawString(100, 710, f"Phone: {resume.phone}")
    p.drawString(100, 690, f"Address: {resume.address}")
    p.drawString(100, 670, f"Nationality: {resume.nationality}")
    p.drawString(100, 650, f"Summary: {resume.summary}")

    # Finish the PDF
    p.showPage()
    p.save()
    buffer.seek(0)

    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{resume.First_Name}_Resume.pdf"'
    return response

# Helper function to generate DOCX
def generate_docx(resume):
    from docx import Document

    doc = Document()
    doc.add_heading(f"{resume.First_Name} {resume.Last_Name}", level=1)
    doc.add_paragraph(f"Email: {resume.email}")
    doc.add_paragraph(f"Phone: {resume.phone}")
    doc.add_paragraph(f"Address: {resume.address}")
    doc.add_paragraph(f"Nationality: {resume.nationality}")
    doc.add_paragraph(f"Summary: {resume.summary}")

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="{resume.first_name}_Resume.docx"'

    doc.save(response)
    return response

# Helper function to generate JPEG
def generate_jpeg(resume):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (800, 600), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 50), f"Name: {resume.First_Name} {resume.Last_Name}", fill="black")
    draw.text((50, 80), f"Email: {resume.email}", fill="black")
    draw.text((50, 110), f"Phone: {resume.phone}", fill="black")
    draw.text((50, 140), f"Address: {resume.address}", fill="black")
    draw.text((50, 170), f"Summary: {resume.summary}", fill="black")

    buffer = BytesIO()
    image.save(buffer, format="JPEG")
    buffer.seek(0)

    response = HttpResponse(buffer, content_type="image/jpeg")
    response["Content-Disposition"] = f'attachment; filename="{resume.first_name}_Resume.jpeg"'
    return response

@login_required
@csrf_exempt  # Use this only for testing; consider using CSRF tokens in production
def generate_pdf(request):
    resume = Resume.objects.filter(user=request.user).first()
    if not resume:
        return HttpResponse("No resume data found.")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=18)
    normal_style = styles["BodyText"]

    content = []
    content.append(Paragraph(f"{resume.First_Name} {resume.Last_Name}", title_style))
    content.append(Paragraph(f"Email: {resume.email}", normal_style))
    content.append(Paragraph(f"Phone: {resume.phone}", normal_style))
    content.append(Paragraph(f"Address: {resume.address}", normal_style))
    content.append(Paragraph(f"Nationality: {resume.nationality}", normal_style))
    content.append(Paragraph("<b>About You:</b>", title_style))
    about_you_text = resume.about_you if resume.about_you else "No description provided."
    content.append(Paragraph(about_you_text, normal_style))
    # Add other fields similarly...

    doc.build(content)

    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="resume.pdf"'
    response.write(pdf)
    return response
